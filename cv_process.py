from docx import Document
import openai
import json
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Pour aligner les paragraphes
from docx.shared import Inches  # Pour définir les positions en pouces
import docx2txt
import streamlit as st
import tempfile
import re
from datetime import datetime
from pydantic import BaseModel, Field
from typing import List
from openai import OpenAI
import fitz  # PyMuPDF


openai.api_key = st.secrets["OPENAI_API_KEY"]

client = OpenAI()


def preprocess_text(text: str) -> str:
    """
    Nettoie le texte extrait avant envoi au LLM.
    
    Args:
        text: Texte brut à nettoyer
        
    Returns:
        Texte nettoyé
    """
    # 1. Supprimer les codes Unicode (\uXXXX)
    text = re.sub(r'\\u[0-9a-fA-F]{4}', '', text)
    
    # 2. Remplacer les espaces multiples par un seul espace
    text = re.sub(r' +', ' ', text)
    
    # 3. Remplacer les sauts de ligne multiples par double saut de ligne
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extrait le texte d'un fichier PDF avec preprocessing.
    
    Args:
        file_path: Chemin vers le fichier PDF
        
    Returns:
        Texte extrait et nettoyé du PDF
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        
        return preprocess_text(text)
        
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction PDF: {str(e)}")


def extract_text_from_word(file_path: str) -> str:
    """
    Extrait le texte d'un fichier Word avec preprocessing.
    
    Args:
        file_path: Chemin vers le fichier Word
        
    Returns:
        Texte extrait et nettoyé du fichier Word
    """
    try:
        document = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        
        return preprocess_text(text)
        
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction Word: {str(e)}")


def extract_text_from_file(file_path: str) -> str:
    """
    Extrait le texte d'un fichier, qu'il soit PDF ou Word.

    Args:
        file_path: Chemin vers le fichier

    Returns:
        Texte extrait du fichier
    """
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_word(file_path)
    else:
        raise ValueError("Format de fichier non supporté. Seuls les fichiers PDF et Word sont acceptés.")

    
def generate_trigramme(prenom, nom):
    """
    Génère le trigramme : première lettre du prénom + deux premières consonnes du nom.
    Exemple : Cédric GOBERT -> CGB
    """
    prenom = prenom.strip().upper() if prenom else ""
    nom = nom.strip().upper() if nom else ""
    # Enlever les espaces pour les noms composés
    nom_sans_espace = nom.replace(" ", "")
    first_letter = prenom[0] if prenom else ""
    consonnes = re.sub(r'[AEIOUY]', '', nom_sans_espace)
    trigramme = first_letter + consonnes[:2]
    return trigramme
    

class Projet(BaseModel):
    CLIENT_NOM: str = Field(..., description="Nom du client.")
    DATE_DEBUT: str = Field(..., description="Date de début du projet au format MM/AAAA.")
    DATE_FIN: str = Field(..., description="Date de fin du projet au format MM/AAAA.")
    INTITULE_POSTE: str = Field(..., description="Intitulé du poste occupé.")
    INTITULE_PROJET: str = Field(..., description="Intitulé du projet réalisé.")
    DETAILS_PROJET: str = Field(None, description="Informations supplémentaires tel que le budget, les effectifs et les heures sans accident")
    REALISATION: List[str] = Field(..., description="Réalisations principales du projet.")


class Diplome(BaseModel):
    ANNEE_DIPLOME: str = Field(..., description="Année d'obtention du diplôme.")
    INTITULE_DIPLOME: str = Field(..., description="Intitulé complet du diplôme obtenu.")


class Langue(BaseModel):
    LANGUE: str = Field(..., description="Nom de la langue parlée.")
    NIVEAU: str = Field(..., description="Niveau de maîtrise de la langue (exemple : Courant, Intermédiaire).")


class FormationComplementaire(BaseModel):
    ANNEE_FORMATION: str = Field(..., description="Année de la formation complémentaire.")
    INTITULE_FORMATION: str = Field(..., description="Intitulé complet de la formation complémentaire.")


class CVInfo(BaseModel):
    PRENOM: str = Field(..., description="prénom")
    NOM: str = Field(..., description="nom")
    AGE: str = Field(..., description="âge")
    INTITULE_DU_POSTE: str = Field(..., description="L'intitulé du poste recherché.")
    EXPERTISE: List[str] = Field(..., description="Les activités et compétences spécifiques (par exemple, Etude de constructibilité, Résolution des problématiques, Leadership).")
    SECTEUR: List[str] = Field(..., description="Les domaines principaux d'expertise (par exemple, Bâtiment, Industrie, Oil & Gas).")
    METHODOLOGIE: List[str] = Field(..., description="Les méthodologies et outils maîtrisés (par exemple, Pack office, MS Project, Naviswork).")
    HABILITATION: List[str] = Field(..., description="Les habilitations professionnelles spécifiques (par exemple, GIES 1/2, Elf Gabon HS3).")
    Projets_effectués: List[Projet] = Field(..., description="Liste des projets effectués avec les détails de chaque mission.")
    Diplômes: List[Diplome] = Field(..., description="Liste des diplômes obtenus.")
    Langues: List[Langue] = Field(..., description="Langues parlées avec leur niveau de maîtrise.")
    Formations_complémentaires: List[FormationComplementaire] = Field(..., description="Formations complémentaires suivies.")


def extract_info_from_cv(cv_text: str, language: str = "fr") -> CVInfo:
    """
    Extrait des informations structurées à partir d'un texte de CV en utilisant l'API OpenAI.
    
    Arguments :
        cv_text (str) : Contenu textuel du CV.

    Retourne :
        CVInfo : Un objet Pydantic contenant les informations extraites.
    """
    system_prompt = {
        "fr": "Tu es un assistant qui aide à extraire les informations des CV.",
        "en": "You are an assistant that helps extract information from resumes. Extract the required fields in english."
    }

    system_prompt = system_prompt.get(language, system_prompt["fr"])
    
    completion = client.chat.completions.parse(
        model="gpt-5",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": cv_text},
        ],
        response_format=CVInfo,  
    )

    parsed: CVInfo = completion.choices[0].message.parsed
    info = parsed.model_dump()

    # Générer le trigramme localement
    prenom = info.get("PRENOM", "")
    nom = info.get("NOM", "")
    info["NOM"] = nom.upper()  # Forcer le nom en majuscule
    info["TRI"] = generate_trigramme(prenom, nom)

    # Extraire l'email via regex sur le texte du CV
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", cv_text)
    if email_match:
        info["EMAIL"] = email_match.group(0)
    else:
        info["EMAIL"] = ""

    # Utiliser la valeur AGE extraite par l'API pour calculer l'année de naissance
    age_str = info.get("AGE", "")
    try:
        age = int(age_str)
        current_year = datetime.now().year
        info["ANNEE"] = current_year - age
    except (ValueError, TypeError):
        info["ANNEE"] = ""

    # Extraire le téléphone via regex sur le texte du CV
    tel_match = re.search(r'(\d{2}(?:[\s\.-]?\d{2}){4})', cv_text)
    if tel_match:
        # Nettoyer le numéro pour enlever espaces, tirets, points
        raw_tel = tel_match.group(1)
        digits = re.sub(r'[^0-9]', '', raw_tel)
        # Reformater en XX.XX.XX.XX.XX
        if len(digits) == 10:
            info["TELEPHONE"] = '.'.join([digits[i:i+2] for i in range(0, 10, 2)])
        else:
            info["TELEPHONE"] = raw_tel
    else:
        info["TELEPHONE"] = ""

    return info


def fill_word_template_with_lists(template_path, output_path, data, language="fr"):
    """
    Remplit un modèle Word avec des données (y compris dans l'en-tête),
    en remplaçant les placeholders et en appliquant les styles nécessaires.

    Arguments :
        template_path (str) : Chemin vers le modèle Word.
        output_path (str) : Chemin vers le fichier Word généré.
        data (dict) : Données à insérer dans le fichier Word.
    """
    doc = Document(template_path)

    # --- 🔹 1. Gestion des en-têtes et pieds de page ---
    for section in doc.sections:
        header = section.header
        footer = section.footer

        # En-tête
        for paragraph in header.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Pied de page (optionnel, même logique)
        for paragraph in footer.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))


    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # Placeholder au format {{KEY}}

            # --- Projets effectués ---
            if key == "Projets_effectués" and isinstance(value, list):
                if placeholder in paragraph.text:
                    paragraph.text = ""
                    for projet in value:
                        client_nom = projet.get('CLIENT_NOM') or 'Non spécifié'
                        dates = f"{(projet.get('DATE_DEBUT') or 'N/A')} - {(projet.get('DATE_FIN') or 'N/A')}"
                        client_date_line = f"{client_nom}\t{dates}"

                        client_date_paragraph = paragraph.insert_paragraph_before(client_date_line)
                        client_date_paragraph.style = "italique gras"

                        tab_stops = client_date_paragraph.paragraph_format.tab_stops
                        tab_stop = tab_stops.add_tab_stop(Inches(6.5))
                        tab_stop.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        post_paragraph = paragraph.insert_paragraph_before(projet.get('INTITULE_POSTE') or 'Non spécifié')
                        post_paragraph.style = paragraph.style

                        paragraph.insert_paragraph_before("")

                        project_paragraph = paragraph.insert_paragraph_before(projet.get('INTITULE_PROJET') or 'Non spécifié')
                        project_paragraph.style = paragraph.style
                        project_paragraph.runs[0].bold = True

                        details_projet = (projet.get('DETAILS_PROJET') or '').strip()
                        if details_projet:
                            project_paragraph = paragraph.insert_paragraph_before(details_projet)
                            project_paragraph.style = paragraph.style

                        paragraph.insert_paragraph_before("")

                        realizations = projet.get('REALISATION') or []
                        if realizations:
                            title_realizations = "Réalisations :" if language == "fr" else "Achievements :"
                            realizations_paragraph = paragraph.insert_paragraph_before(title_realizations)
                            realizations_paragraph.style = paragraph.style
                            realizations_paragraph.runs[0].bold = True

                            for realization in realizations:
                                realization = (realization or '').strip()
                                if realization:
                                    realization_paragraph = paragraph.insert_paragraph_before(realization)
                                    realization_paragraph.style = "Liste à puces1"

                            paragraph.insert_paragraph_before("")

            # --- Diplômes ---
            elif key == "Diplômes" and isinstance(value, list):
                if placeholder in paragraph.text:
                    paragraph.text = ""
                    for diplome in value:
                        diploma_line = f"{(diplome.get('ANNEE_DIPLOME') or 'N/A')}    {(diplome.get('INTITULE_DIPLOME') or 'Non spécifié')}"
                        diploma_paragraph = paragraph.insert_paragraph_before(diploma_line)
                        diploma_paragraph.style = paragraph.style

            # --- Langues ---
            elif key == "Langues" and isinstance(value, list):
                if placeholder in paragraph.text:
                    paragraph.text = ""
                    for langue in value:
                        language_line = f"{(langue.get('LANGUE') or 'Non spécifié')}    {(langue.get('NIVEAU') or 'Non spécifié')}"
                        language_paragraph = paragraph.insert_paragraph_before(language_line)
                        language_paragraph.style = paragraph.style

            # --- Formations complémentaires ---
            elif key == "Formations_complémentaires" and isinstance(value, list):
                if placeholder in paragraph.text:
                    paragraph.text = ""
                    for formation in value:
                        formation_line = f"{(formation.get('ANNEE_FORMATION') or 'N/A')}    {(formation.get('INTITULE_FORMATION') or 'Non spécifié')}"
                        formation_paragraph = paragraph.insert_paragraph_before(formation_line)
                        formation_paragraph.style = paragraph.style

            # --- Listes génériques ---
            elif isinstance(value, list):
                if placeholder in paragraph.text:
                    paragraph.text = ""
                    for item in value:
                        list_paragraph = paragraph.insert_paragraph_before(str(item))
                        list_paragraph.style = paragraph.style

            # --- Valeurs simples ---
            elif placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # --- 🔹 3. Sauvegarde du fichier final ---
    doc.save(output_path)
