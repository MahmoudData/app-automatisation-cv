from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Couleur bleue pour tous les textes issus des données
BLUE = RGBColor(0x1F, 0x48, 0x7C)
BLACK = RGBColor(0x00, 0x00, 0x00)


def replace_in_paragraph_inter(paragraph, placeholder, value, use_black=False):
    """
    Remplace un placeholder dans un paragraphe en préservant le formatage du texte existant
    et en appliquant la couleur spécifique uniquement sur la valeur remplacée.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return

    # Trouver la position du placeholder
    placeholder_start = full_text.find(placeholder)
    placeholder_end = placeholder_start + len(placeholder)
    
    value_str = str(value)
    new_text = full_text.replace(placeholder, value_str, 1)
    
    # Reconstruire le paragraphe
    current_pos = 0
    new_runs = []
    
    for run in paragraph.runs:
        run_len = len(run.text)
        run_end = current_pos + run_len
        
        if run_end <= placeholder_start:
            # Ce run est avant le placeholder, on le garde tel quel
            new_runs.append((run.text, run, False))
        elif current_pos >= placeholder_end:
            # Ce run est après le placeholder, on le garde tel quel
            adjusted_pos = current_pos - len(placeholder) + len(value_str)
            new_runs.append((new_text[adjusted_pos:adjusted_pos + run_len], run, False))
        else:
            # Ce run contient une partie du placeholder
            if current_pos < placeholder_start:
                # Partie avant le placeholder
                before_len = placeholder_start - current_pos
                new_runs.append((run.text[:before_len], run, False))
            
            # La valeur de remplacement (seulement si c'est le premier run qui la contient)
            if placeholder_start >= current_pos and placeholder_start < run_end:
                new_runs.append((value_str, run, True))  # True = appliquer la couleur spéciale
            
            # Partie après le placeholder dans ce run
            if run_end > placeholder_end:
                after_start = placeholder_end - current_pos
                new_runs.append((run.text[after_start:], run, False))
        
        current_pos = run_end
    
    # Vider tous les runs
    for run in paragraph.runs:
        run.text = ""
    
    # Recréer les runs avec le bon formatage
    run_idx = 0
    for text, original_run, apply_color in new_runs:
        if text:
            if run_idx < len(paragraph.runs):
                target_run = paragraph.runs[run_idx]
            else:
                target_run = paragraph.add_run()
            
            target_run.text = text
            
            # Copier le formatage de l'original
            if original_run.bold is not None:
                target_run.bold = original_run.bold
            if original_run.italic is not None:
                target_run.italic = original_run.italic
            if original_run.font.size is not None:
                target_run.font.size = original_run.font.size
            
            # Appliquer la couleur uniquement sur la valeur remplacée
            if apply_color:
                target_run.font.color.rgb = BLACK if use_black else BLUE
            elif original_run.font.color.rgb is not None:
                target_run.font.color.rgb = original_run.font.color.rgb
            
            run_idx += 1


def normalize_key(key):
    """Normalise une clé pour la recherche insensible à la casse."""
    return key.lower().replace('_', '').replace('-', '')


def find_matching_key(placeholder_key, data):
    """Trouve la clé correspondante dans les données, insensible à la casse."""
    normalized_placeholder = normalize_key(placeholder_key)
    
    for key, value in data.items():
        if normalize_key(key) == normalized_placeholder:
            return value
    
    return None


def fill_cv_inter_template(template_path, output_path, data, language="fr"):
    """
    Version finale qui gère tous les cas de placeholders.
    """
    doc = Document(template_path)
    
    # Liste des clés à mettre en noir (poste_2 est en noir, poste_1 est en bleu)
    black_keys = {
        'poste2', 'nomemploye', 'education', 'affiliations', 
        'datenaissance', 'nombreannees', 'nationalite', 'attributions'
    }

    # --- 🔹 1. Remplacement dans les en-têtes ---
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            text = paragraph.text
            placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
            
            for placeholder_key in placeholders:
                value = find_matching_key(placeholder_key, data)
                if value is not None and not isinstance(value, (list, dict)):
                    placeholder = f"{{{{{placeholder_key}}}}}"
                    display_value = str(value) if value not in [None, ""] else ""
                    # Vérifier si cette clé doit être en noir
                    use_black = normalize_key(placeholder_key) in black_keys
                    replace_in_paragraph_inter(paragraph, placeholder, display_value, use_black)

    # --- 🔹 2. Remplacement dans les paragraphes ---
    paragraphs_to_remove = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        placeholders = re.findall(r'\{\{([^}]+)\}\}', original_text)
        
        if not placeholders:
            continue
            
        # Flag pour savoir si le paragraphe a été traité
        paragraph_handled = False
        
        for placeholder_key in placeholders:
            placeholder = f"{{{{{placeholder_key}}}}}"
            value = find_matching_key(placeholder_key, data)
            
            normalized = normalize_key(placeholder_key)
            
            # --- 🔸 QUALIFICATIONS ---
            if normalized == normalize_key("QUALIFICATIONS"):
                qualifications = find_matching_key("QUALIFICATIONS", data)
                if qualifications and isinstance(qualifications, list) and len(qualifications) > 0:
                    paragraph.clear()
                    for qualification in qualifications:
                        qual_paragraph = paragraph.insert_paragraph_before(str(qualification))
                        qual_paragraph.style = "Liste à puces1"
                        # couleur bleue sur les textes provenant des données
                        for run in qual_paragraph.runs:
                            run.font.color.rgb = BLUE
                    paragraph_handled = True
                    break
                else:
                    replace_in_paragraph_inter(paragraph, placeholder, "")
                    paragraph_handled = True
                    break

            # --- 🔸 DIPLOMES (détection de ligne template) ---
            elif normalized in ['diplomes', 'annee', 'diplome', 'ecole', 'ville']:
                text_lower = original_text.lower()
                # Vérifier si c'est une ligne template complète
                if ("annee" in text_lower or "ANNEE" in original_text) and \
                ("diplome" in text_lower or "DIPLOME" in original_text):
                    diplomes = find_matching_key("DIPLOMES", data)
                    if diplomes and isinstance(diplomes, list):
                        # Supprimer le contenu du paragraphe template
                        paragraph.clear()
                        
                        for idx, diplome in enumerate(diplomes):
                            annee = diplome.get('ANNEE', 'N/A')
                            diplome_nom = diplome.get('DIPLOME', 'Non spécifié')
                            ecole = diplome.get('ECOLE', 'Non spécifié')
                            ville = diplome.get('VILLE', 'Non spécifié')
                            
                            # Créer un nouveau paragraphe pour chaque diplôme
                            new_para = paragraph.insert_paragraph_before()
                            new_para.style = paragraph.style
                            
                            # Ajouter l'année
                            run_annee = new_para.add_run(str(annee))
                            
                            # Ajouter une tabulation
                            new_para.add_run('\t')
                            
                            # Ajouter le reste (tout en bloc)
                            diplome_text = f"{diplome_nom} – {ecole} – {ville}"
                            run_diplome = new_para.add_run(diplome_text)
                            
                            # APPLIQUER LE HANGING INDENT + MARGE DROITE 
                            new_para.paragraph_format.left_indent = Cm(3.5)  # Marge gauche
                            new_para.paragraph_format.right_indent = Cm(1.5)  # Marge droite identique
                            new_para.paragraph_format.first_line_indent = Cm(-2.5)  # Retrait négatif pour la date
                            
                            # AJOUTER UN ESPACE APRÈS CHAQUE DIPLÔME (sauf le dernier) 
                            if idx < len(diplomes) - 1:
                                new_para.paragraph_format.space_after = Pt(12)
                            
                            # Optionnel : définir la taille de police
                            run_annee.font.size = Pt(10)
                            run_diplome.font.size = Pt(10)

                            # couleur bleue sur les textes de diplôme
                            run_annee.font.color.rgb = BLUE
                            run_diplome.font.color.rgb = BLUE
                        
                        paragraph_handled = True
                        break

            # --- 🔸 CERTIFICATIONS ---
            elif normalized == normalize_key("CERTIFICATIONS"):
                certifications = find_matching_key("CERTIFICATIONS", data)
                if certifications and isinstance(certifications, list) and len(certifications) > 0:
                    paragraph.clear()
                    for certification in certifications:
                        cert_paragraph = paragraph.insert_paragraph_before(str(certification))
                        cert_paragraph.style = "Liste à puces1"
                        # couleur bleue sur les textes provenant des données
                        for run in cert_paragraph.runs:
                            run.font.color.rgb = BLUE
                    paragraph_handled = True
                    break
                else:
                    replace_in_paragraph_inter(paragraph, placeholder, "")
                    paragraph_handled = True
                    break

            # --- 🔸 PROJETS ---
            elif normalized == normalize_key("PROJETS"):
                projets = find_matching_key("PROJETS", data)
                if projets and isinstance(projets, list) and len(projets) > 0:
                    paragraph.clear()
                    for idx_projet, projet in enumerate(projets):
                        client_nom = projet.get('CLIENT_NOM', '')
                        dates = projet.get('DATES', '')
                        intitule_poste = projet.get('INTITULE_POSTE', '')
                        intitule_projet = projet.get('INTITULE_PROJET', '')
                        realisations = projet.get('REALISATION', [])

                        # Ligne 1 : CLIENT_NOM (TAB) DATES (en gras)
                        para_client = paragraph.insert_paragraph_before()
                        if client_nom:
                            run_client = para_client.add_run(client_nom)
                            run_client.bold = True
                            run_client.font.color.rgb = BLUE
                            para_client.add_run('\t')
                        run_dates = para_client.add_run(dates)
                        run_dates.bold = True
                        run_dates.font.color.rgb = BLUE
                        para_client.paragraph_format.left_indent = Cm(0.5)
                        para_client.paragraph_format.right_indent = Cm(0.5)


                        # Ligne 2 : INTITULE_POSTE (en gras)
                        para_poste = paragraph.insert_paragraph_before()
                        run_poste = para_poste.add_run(intitule_poste)
                        run_poste.bold = True
                        run_poste.font.color.rgb = BLUE
                        para_poste.paragraph_format.left_indent = Cm(0.5)
                        para_poste.paragraph_format.right_indent = Cm(0.5)

                        # Saut de ligne
                        paragraph.insert_paragraph_before()

                        # Ligne 3 : INTITULE_PROJET
                        para_intitule = paragraph.insert_paragraph_before(intitule_projet)
                        for run in para_intitule.runs:
                            run.font.color.rgb = BLUE
                        para_intitule.paragraph_format.left_indent = Cm(0.5)
                        para_intitule.paragraph_format.right_indent = Cm(0.5)

                        # Saut de ligne
                        paragraph.insert_paragraph_before()

                        # Réalisations (une par ligne)
                        for realisation in realisations:
                            para_real = paragraph.insert_paragraph_before(str(realisation))
                            for run in para_real.runs:
                                run.font.color.rgb = BLUE
                            para_real.paragraph_format.left_indent = Cm(0.5)
                            para_real.paragraph_format.right_indent = Cm(0.5)

                        # Ajouter un espace entre les projets (sauf le dernier)
                        if idx_projet < len(projets) - 1:
                            para_espace = paragraph.insert_paragraph_before()
                            para_espace.paragraph_format.space_after = Pt(12)

                    paragraph_handled = True
                    break
                else:
                    replace_in_paragraph_inter(paragraph, placeholder, "")
                    paragraph_handled = True
                    break

            # --- 🔸 PAYS ---
            elif normalized == normalize_key("PAYS"):
                if value is not None and not isinstance(value, (list, dict)):
                    replace_in_paragraph_inter(paragraph, placeholder, str(value))
                    paragraph_handled = True
                    break

            # --- 🔸 Valeurs simples ---
            elif value is not None and not isinstance(value, (list, dict)):
                display_value = str(value) if value not in [None, ""] else ""
                # Vérifier si cette clé doit être en noir
                use_black = normalize_key(placeholder_key) in black_keys
                replace_in_paragraph_inter(paragraph, placeholder, display_value, use_black)
                paragraph_handled = True
            
            # --- 🔸 Valeurs NULL ou vides ---
            elif value is None or value == "":
                # Remplacer par chaîne vide
                replace_in_paragraph_inter(paragraph, placeholder, "")
                paragraph_handled = True

    # --- 🔹 3. Remplacement dans les tableaux ---
    for table in doc.tables:
        
        # --- 🔸 Tableau des INFORMATIONS GÉNÉRALES (Tableau 1) ---
        if len(table.columns) == 2 and len(table.rows) == 9:
            # C'est le tableau avec Poste, Nom du consultant, etc.
            for row in table.rows:
                cells = row.cells
                # La colonne 2 (index 1) contient les placeholders
                if len(cells) >= 2:
                    for paragraph in cells[1].paragraphs:
                        text = paragraph.text
                        placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
                        
                        for placeholder_key in placeholders:
                            value = find_matching_key(placeholder_key, data)
                            placeholder = f"{{{{{placeholder_key}}}}}"
                            
                            # Remplacer par la valeur OU par vide si None/vide
                            if value is not None and not isinstance(value, (list, dict)):
                                display_value = str(value) if value != "" else ""
                            else:
                                display_value = ""  # Remplacer par vide si pas de valeur
                            
                            # Vérifier si cette clé doit être en noir
                            use_black = normalize_key(placeholder_key) in black_keys
                            replace_in_paragraph_inter(paragraph, placeholder, display_value, use_black)
    
        # --- 🔸 Tableau des EXPÉRIENCES (utilise maintenant la liste PROJETS unifiée) ---
        if len(table.columns) == 4 and any("Période" in cell.text or "Period" in cell.text for cell in table.rows[0].cells):
            rows_to_remove = []
            for idx, row in enumerate(table.rows[1:], start=1):
                cell_texts = [cell.text for cell in row.cells]
                if any("{{" in text for text in cell_texts):
                    rows_to_remove.append(idx)

            for idx in sorted(rows_to_remove, reverse=True):
                tbl = table._element
                tr = table.rows[idx]._element
                tbl.remove(tr)

            # Utiliser la liste PROJETS unifiée
            projets = find_matching_key("PROJETS", data)
            if projets:
                for projet in projets:
                    row_cells = table.add_row().cells
                    # Colonne 1 : DATES (période du projet)
                    row_cells[0].text = projet.get('DATES', 'N/A')
                    # Colonne 2 : ENTREPRISE
                    row_cells[1].text = projet.get('ENTREPRISE', 'Non spécifié')
                    # Colonne 3 : INTITULE_POSTE (poste précis)
                    row_cells[2].text = projet.get('INTITULE_POSTE', 'Non spécifié')

                    # Colonne 4 : Chantiers suivis (format CLIENT_NOM : INTITULE_PROJET)
                    cell_experience = row_cells[3]
                    cell_experience._element.clear_content()

                    client_nom = projet.get('CLIENT_NOM', '')
                    intitule_projet = projet.get('INTITULE_PROJET', '')

                    # Format : CLIENT_NOM : INTITULE_PROJET
                    if client_nom:
                        chantier_text = f"{client_nom} : {intitule_projet}"
                    else:
                        chantier_text = intitule_projet

                    para_chantier = cell_experience.add_paragraph(chantier_text)
                    for run in para_chantier.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = BLUE
                    para_chantier.paragraph_format.left_indent = Cm(0.5)
                    para_chantier.paragraph_format.right_indent = Cm(0.5)

                    # Centrer les colonnes Dates, Entreprise, Fonction
                    for idx in [0, 1, 2]:
                        for para in row_cells[idx].paragraphs:
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in para.runs:
                                run.font.size = Pt(10)
                                run.font.color.rgb = BLUE

        # --- 🔸 Tableau des LANGUES ---
        elif len(table.columns) == 4 and any("Langue" in cell.text or "Language" in cell.text for cell in table.rows[0].cells):

            langues = find_matching_key("LANGUES", data)

            # Supprimer les lignes template
            rows_to_remove = []
            for idx, row in enumerate(table.rows[1:], start=1):
                if any("{{" in cell.text for cell in row.cells):
                    rows_to_remove.append(idx)

            for idx in sorted(rows_to_remove, reverse=True):
                tbl = table._element
                tr = table.rows[idx]._element
                tbl.remove(tr)

            # Ajouter les nouvelles lignes
            if langues:
                for langue in langues:
                    row_cells = table.add_row().cells

                    row_cells[0].text = langue.get('LANGUE', 'Non spécifié')
                    comp = langue.get('COMPETENCE', {})
                    row_cells[1].text = comp.get('LUE', 'N/A')
                    row_cells[2].text = comp.get('ECRITE', 'N/A')
                    row_cells[3].text = comp.get('PARLEE', 'N/A')

                    # Appliquer le fond GRIS directement (D9D9D9)
                    for cell in row_cells:
                        shading = parse_xml(
                            r'<w:shd {} w:fill="F1F1F1"/>'.format(nsdecls('w'))
                        )
                        cell._tc.get_or_add_tcPr().append(shading)

                    # Texte en noir
                    for cell in row_cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # NOIR

    doc.save(output_path)
    print(f"✅ Document généré avec succès : {output_path}")
